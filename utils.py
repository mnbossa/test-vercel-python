import re
from docx import Document
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph
import json
import pandas as pd
# from docx.document import Document as _Document

def parse_xml_tags(text):
    """Parse XML-like tags from text into a dictionary"""
    tags = {}
    tag_pattern = re.compile(r'<([^>]+)>(.*?)</\1>', re.DOTALL)
    for match in tag_pattern.finditer(text):
        tag_name = match.group(1).split(' ')[0].strip('<>')  # Handle tags with attributes
        content = match.group(2).strip()
        # Handle multiple entries with same tag
        if tag_name in tags:
            if not isinstance(tags[tag_name], list):
                tags[tag_name] = [tags[tag_name]]
            tags[tag_name].append(content)
        else:
            tags[tag_name] = content
    return tags

def iter_block_items(parent):
    """Iterate through document elements in order (paragraphs and tables)"""
    for child in parent.element.body.iterchildren():
        if isinstance(child, CT_P):
            yield Paragraph(child, parent), 'paragraph'
        elif isinstance(child, CT_Tbl):
            yield Table(child, parent), 'table'

def process_table_cell(cell, column_type):
    """Process text in table cells with column context"""
    formatted = []
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            text = run.text.strip()
            if not text:
                continue

            # Check bold AND italic formatting
            if run.font.bold and run.font.italic:
                if column_type == "original":
                    formatted.append(f"[DEL]{text}[/DEL]")  # Deletions in original column
                elif column_type == "amended":
                    formatted.append(f"[ADD]{text}[/ADD]")  # Additions in amended column
            else:
                formatted.append(text)
    return " ".join(formatted)

def parse_amendments(docx_path):
    """Parse document into structured data with header and amendments"""
    doc = Document(docx_path)
    
    header_data = []
    amendments = {}
    current_amendment = None
    collecting_header = True

    for element, elem_type in iter_block_items(doc):
        if elem_type == 'paragraph':
            text = element.text.strip()
            
            # Check for amendment header
            amendment_match = re.search(r'<NumAm>(\d{1,4})</NumAm>', text)
            if amendment_match:
                collecting_header = False
                amendment_num = int(amendment_match.group(1))
                current_amendment = {
                    'raw_content': [],
                }
                amendments[amendment_num] = current_amendment
                continue
                
            if collecting_header:
                header_data.append(text)
            elif current_amendment is not None:
                current_amendment['raw_content'].append(text)
                
        elif elem_type == 'table' and current_amendment is not None:
            # current_amendment[process_table_cell(element.columns[0].cells[1], '')] = "".join([ process_table_cell(cell, 'original')  for cell in element.columns[0].cells[2:]])
            # current_amendment[process_table_cell(element.columns[1].cells[1], '')] = "".join([ process_table_cell(cell, 'amended')  for cell in element.columns[1].cells[2:]])
            if len(element.columns) > 1:
              if len(element.columns[0].cells) > 2:
                current_amendment['OriginalType'] = process_table_cell(element.columns[0].cells[1], '')
                current_amendment['Original'] = '\n'.join([process_table_cell(c, 'original') for c in element.columns[0].cells[2:]])
                current_amendment['AmendedType'] = process_table_cell(element.columns[1].cells[1], '')
                current_amendment['Amended'] = '\n'.join([process_table_cell(c, 'amended') for c in element.columns[1].cells[2:]]) 

    # Process header
    header_text = '\n'.join(header_data)
    parsed_header = parse_xml_tags(header_text)

    # Process amendments
    parsed_amendments = {}
    for num, data in amendments.items():
        full_content = '\n'.join(data['raw_content'])
        amendment_data = parse_xml_tags(full_content)
        
        # Add amendment changes
        for k, v in data.items():
           amendment_data[k] = v
        
        # Add raw content for fallback
        amendment_data['raw_content'] = full_content

        members_str = amendment_data['RepeatBlock-By']
        members_list = [
            name.strip()
            for name in members_str.replace('<Members>', '')
                                  .replace('</Members>', '')
                                  .split(',')
        ]
        amendment_data['By'] = members_list
        del amendment_data['RepeatBlock-By']

        if 'raw_content' in amendment_data:
            justification_match = re.search(
                r'<TitreJust>Justification</TitreJust>(.*?)(<|$)',
                amendment_data['raw_content'],
                re.DOTALL
            )
            if justification_match:
                justification_text = justification_match.group(1).strip()
                amendment_data['Justification'] = justification_text
 
        parsed_amendments[num] = amendment_data

    return {
        'header': parsed_header,
        'amendments': parsed_amendments
    }



def parse_draft_report(docx_path):
    """Parse document into structured data with header and amendments"""
    doc = Document(docx_path)
    
    header_data = []
    amendments = {}
    current_amendment = None
    collecting_header = True

    for element, elem_type in iter_block_items(doc):
        if elem_type == 'paragraph':
            text = element.text.strip()
            
            # Check for amendment header
            amendment_match = re.search(r'<NumAm>(\d{1,4})</NumAm>', text)
            if amendment_match:
                collecting_header = False
                amendment_num = int(amendment_match.group(1))
                current_amendment = {
                    'raw_content': [],
                }
                amendments[amendment_num] = current_amendment
                continue
                
            if collecting_header:
                header_data.append(text)
            elif current_amendment is not None:
                current_amendment['raw_content'].append(text)
                
        elif elem_type == 'table' and current_amendment is not None:
            # current_amendment[process_table_cell(element.columns[0].cells[1], '')] = "".join([ process_table_cell(cell, 'original')  for cell in element.columns[0].cells[2:]])
            # current_amendment[process_table_cell(element.columns[1].cells[1], '')] = "".join([ process_table_cell(cell, 'amended')  for cell in element.columns[1].cells[2:]])
            if len(element.columns) > 1:
              if len(element.columns[0].cells) > 2:
                current_amendment['OriginalType'] = process_table_cell(element.columns[0].cells[1], '')
                current_amendment['Original'] = '\n'.join([process_table_cell(c, 'original') for c in element.columns[0].cells[2:]])
                current_amendment['AmendedType'] = process_table_cell(element.columns[1].cells[1], '')
                current_amendment['Amended'] = '\n'.join([process_table_cell(c, 'amended') for c in element.columns[1].cells[2:]]) 

    # Process header
    header_text = '\n'.join(header_data)
    parsed_header = parse_xml_tags(header_text)

    # Process amendments
    parsed_amendments = {}
    for num, data in amendments.items():
        full_content = '\n'.join(data['raw_content'])
        amendment_data = parse_xml_tags(full_content)
        
        # Add amendment changes
        for k, v in data.items():
           amendment_data[k] = v
        
        # Add raw content for fallback
        amendment_data['raw_content'] = full_content

 

        if 'raw_content' in amendment_data:
            justification_match = re.search(
                r'<TitreJust>Justification</TitreJust>(.*?)(<|$)',
                amendment_data['raw_content'],
                re.DOTALL
            )
            if justification_match:
                justification_text = justification_match.group(1).strip()
                amendment_data['Justification'] = justification_text
 
        parsed_amendments[num] = amendment_data

    return {
        'header': parsed_header,
        'amendments': parsed_amendments
    }



def get_political_group_abbreviation(full_name):
    """Converts full political group names to their common abbreviations."""
    mapping = {
        "Group of the European People's Party (Christian Democrats)": "EPP",
        "Group of the Progressive Alliance of Socialists and Democrats in the European Parliament": "S&D",
        "Renew Europe Group": "Renew",
        "Group of the Greens/European Free Alliance": "Greens/EFA",
        "European Conservatives and Reformists Group": "ECR",
        "Identity and Democracy Group": "ID",
        "The Left group in the European Parliament - GUE/NGL": "The Left", # Or GUE/NGL
        "Non-attached Members": "NI" # Non-Inscrits
        # Add any other groups if present in your data
    }
    return mapping.get(full_name, full_name) # Return abbreviation or original if not found

def get_spanish_party_abbreviation(full_name):
    """Converts full Spanish national party names to their common abbreviations."""
    mapping = {
        "Partido Popular": "PP",
        "Partido Socialista Obrero Español": "PSOE",
        "VOX": "VOX", # Already an abbreviation, but good to have for consistency
        "Ciudadanos - Partido de la Ciudadanía": "Cs",
        " Unidas Podemos Cambiar Europa": "Podemos", # Or UP, check your data for exact string
        "PODEMOS": "Podemos",
        "IZQUIERDA UNIDA": "IU",
        "SUMAR": "Sumar",
        "Esquerra Republicana de Catalunya": "ERC",
        "Junts per Catalunya": "Junts", # Or JxCat
        "Partido Nacionalista Vasco": "PNV", # EAJ-PNV
        "Euskal Herria Bildu": "EH Bildu",
        "Bloque Nacionalista Galego": "BNG",
        "Coalición Canaria": "CCa",
        # Add other Spanish parties and their full names as they appear in your JSON
    }
    # Handle cases where the mapping might be for a part of the name or needs flexibility
    for key, value in mapping.items():
        if key.lower() in full_name.lower(): # Case-insensitive partial match
            return value
    return full_name # Return original if no specific abbreviation is found

def process_meps_data(input_filepath='../data/meps.json'):
    """
    Reads MEP data, abbreviates political groups and Spanish national parties,
    and saves the processed data.
    """
    try:
        with open(input_filepath, 'r', encoding='utf-8') as f:
            meps_data = json.load(f)
    except FileNotFoundError:
        print(f"Error: Input file '{input_filepath}' not found.")
        return
    except json.JSONDecodeError:
        print(f"Error: Could not decode JSON from '{input_filepath}'.")
        return

    processed_meps = []
    for mep in meps_data:
        processed_mep = mep.copy() # Work on a copy

        # Abbreviate political group
        if 'political_group' in processed_mep:
            processed_mep['political_group'] = get_political_group_abbreviation(processed_mep['political_group'])

        # Abbreviate national party if country is Spain
        if 'country' in processed_mep and processed_mep['country'] == 'Spain':
            if 'national_party' in processed_mep:
                processed_mep['national_party'] = get_spanish_party_abbreviation(processed_mep['national_party'])
        
        processed_meps.append(processed_mep)

    return  processed_meps 

meps = process_meps_data()

def create_amendment_report(result_data):
    """
    Processes MEP and amendment data to create an Excel report.

    Args:
        result_data (dict): A dictionary containing an 'amendments' key.
                            result_data['amendments'] is a dictionary where keys are
                            amendment IDs and values are dictionaries containing a 'By'
                            key (a list of author names).
    """

    # --- 1. Preprocess MEP data for efficient lookup ---
    # Create a lookup dictionary with lowercase names for robust matching
    meps_lookup = {}
    for mep in meps:
        if 'name' in mep and mep['name']: # Ensure name exists and is not empty
            meps_lookup[mep['name'].lower().strip()] = mep
        else:
            print(f"Warning: MEP record found with missing or empty name: {mep}")


    processed_amendments = []

    if 'amendments' not in result_data or not isinstance(result_data['amendments'], dict):
        print("Error: 'amendments' key not found in result_data or is not a dictionary.")
        # Create an empty DataFrame if no amendments to process
        df = pd.DataFrame(columns=["ENMIENDA", "GRUPO", "SPAIN", "Part"])
        return

    # --- 2. Iterate through amendments ---
    for amendment_id, amendment_details in result_data['amendments'].items():
        if 'By' not in amendment_details or not isinstance(amendment_details['By'], list):
            print(f"Warning: Amendment {amendment_id} has no 'By' field or it's not a list. Skipping.")
            # Add a row with empty GRUPO and SPAIN for this amendment if you want to include all amendment IDs
            # processed_amendments.append({
            #     "ENMIENDA": amendment_id,
            #     "GRUPO": "",
            #     "SPAIN": ""
            # })
            continue

        authors = amendment_details['By']
        political_groups_for_amendment = set()
        spanish_national_parties_for_amendment = set()
        found_meps_for_amendment = False

        Part = amendment_details['Article']

        for author_name_raw in authors:
            if not isinstance(author_name_raw, str) or not author_name_raw.strip():
                print(f"Warning: Invalid or empty author name found in amendment {amendment_id}: '{author_name_raw}'")
                continue

            # --- MODIFICATION: Handle names with newline characters ---
            # Take only the part before the first newline character, if present
            cleaned_author_name = author_name_raw.split('\n', 1)[0]
            
            normalized_author_name = cleaned_author_name.lower().strip()
            
            if not normalized_author_name: # Check if name is empty after cleaning
                print(f"Warning: Author name became empty after cleaning in amendment {amendment_id}: Original '{author_name_raw}'")
                continue

            # --- 3. Match author with MEP data ---
            matched_mep = meps_lookup.get(normalized_author_name)

            if matched_mep:
                found_meps_for_amendment = True
                # Extract political group
                if 'political_group' in matched_mep and matched_mep['political_group']:
                    political_groups_for_amendment.add(matched_mep['political_group'])
                
                # Check for Spanish MEP and extract national party
                if 'country' in matched_mep and matched_mep['country'] == 'Spain':
                    if 'national_party' in matched_mep and matched_mep['national_party']:
                        spanish_national_parties_for_amendment.add(matched_mep['national_party'])
            else:
                print(f"Error: MEP not found for author '{normalized_author_name}' in amendment {amendment_id}.")

        # --- 4. Compile results for the current amendment ---
        # Join unique groups and parties into comma-separated strings
        grupo_str = ", ".join(sorted(list(political_groups_for_amendment)))
        spain_str = ", ".join(sorted(list(spanish_national_parties_for_amendment)))
        
        processed_amendments.append({
            "ENMIENDA": amendment_id,
            "GRUPO": grupo_str,
            "SPAIN": spain_str,
            "Part": Part
        })

    # --- 5. Create Pandas DataFrame ---
    df = pd.DataFrame(processed_amendments)

    # Ensure all columns are present even if processed_amendments is empty
    if not processed_amendments:
        df = pd.DataFrame(columns=["ENMIENDA", "GRUPO", "SPAIN", "Part"])

    return df


def create_draft_report(result_data):
    """
    Processes draft report data to create an Excel report.

    Args:
        result_data (dict): A dictionary containing an 'amendments' key.
                            result_data['amendments'] is a dictionary where keys are
                            amendment IDs and values are dictionaries containing a 'By'
                            key (a list of author names).
    """


    processed_amendments = []

    if 'amendments' not in result_data or not isinstance(result_data['amendments'], dict):
        print("Error: 'amendments' key not found in result_data or is not a dictionary.")
        # Create an empty DataFrame if no amendments to process
        df = pd.DataFrame(columns=["ENMIENDA", "GRUPO", "SPAIN", "Part"])
        return

    # --- 2. Iterate through amendments ---
    for amendment_id, amendment_details in result_data['amendments'].items():

        Part = amendment_details['Article']

        
        processed_amendments.append({
            "ENMIENDA": amendment_id,
            "GRUPO": '',
            "SPAIN": '',
            "Part": Part
        })

    # --- 5. Create Pandas DataFrame ---
    df = pd.DataFrame(processed_amendments)

    # Ensure all columns are present even if processed_amendments is empty
    if not processed_amendments:
        df = pd.DataFrame(columns=["ENMIENDA", "GRUPO", "SPAIN", "Part"])

    return df


def save_report(df, output_filename="amendments_analysis.xlsx"):
    # --- 6. Export to Excel ---
    try:
        df.to_excel(output_filename, index=False, engine='openpyxl') # Added engine for broader compatibility
        print(f"Successfully created Excel file: {output_filename}")
    except ImportError:
        print("Error: 'openpyxl' library is required to write Excel files. Please install it using 'pip install openpyxl'.")
        print("Attempting to save as CSV instead.")
        csv_filename = output_filename.replace('.xlsx', '.csv')
        try:
            df.to_csv(csv_filename, index=False)
            print(f"Successfully created CSV file: {csv_filename}")
        except Exception as e_csv:
            print(f"Error writing CSV file: {e_csv}")
    except Exception as e:
        print(f"Error writing Excel file: {e}")


# model = Llama(
#      model_path = "mistral-7b-instruct-v0.2.Q8_0.gguf", # "mistral-7b-instruct-v0.2.Q5_K_M.gguf", # "mistral-7b-instruct-v0.2.Q8_0.gguf",
#      n_ctx=2048,
#      n_gpu_layers=30,  # Offload layers to Metal GPU
#      n_threads=6,     # Use CPU cores
# )

def extract_deletions(text):
    return " | ".join(re.findall(r'\[DEL\](.*?)\[\/DEL\]', text))

def extract_additions(text):
    return " | ".join(re.findall(r'\[ADD\](.*?)\[\/ADD\]', text))

def remove_unnec_tags(text):
    text = re.sub(r'\[/DEL\]\s+\[DEL\]', '', text)
    return re.sub(r'\[/ADD\]\s+\[ADD\]', '', text)


